"""
update_portfolio.py
====================
Reads your resume and updates data.js automatically.
Extracts embedded hyperlinks from PDF/DOCX (LinkedIn, GitHub, Medium, LeetCode, etc.)

Supports: PDF, DOCX, TXT

Usage:
    python update_portfolio.py resume.pdf
    python update_portfolio.py resume.docx
    python update_portfolio.py resume.txt

Requirements:
    pip install pdfplumber python-docx
"""

import sys
import os
import re
import json

# ── 1. Read resume text ───────────────────────────────────────────────────────

def read_pdf(path):
    try:
        import pdfplumber
    except ImportError:
        sys.exit("Run:  pip install pdfplumber")
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += (page.extract_text() or "") + "\n"
    return text.strip()

def read_docx(path):
    try:
        from docx import Document
    except ImportError:
        sys.exit("Run:  pip install python-docx")
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs).strip()

def read_txt(path):
    with open(path, encoding="utf-8") as f:
        return f.read().strip()

def read_resume(path):
    ext = os.path.splitext(path)[1].lower()
    if   ext == ".pdf":  return read_pdf(path)
    elif ext == ".docx": return read_docx(path)
    elif ext == ".txt":  return read_txt(path)
    else: sys.exit(f"Unsupported file type '{ext}'. Use .pdf, .docx, or .txt")

# ── 2. Extract embedded hyperlinks ───────────────────────────────────────────

def extract_pdf_links(path):
    """Extract all URLs embedded as hyperlinks in a PDF."""
    links = []

    # Try pdfplumber annotations
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                # pdfplumber >= 0.7 exposes page.hyperlinks
                if hasattr(page, "hyperlinks"):
                    for h in page.hyperlinks:
                        uri = h.get("uri", "")
                        if uri:
                            links.append(uri)
                # fallback: raw annots dict
                if not links and hasattr(page, "annots") and page.annots:
                    for annot in page.annots:
                        uri = annot.get("uri", "")
                        if uri:
                            links.append(uri)
    except Exception:
        pass

    # Fallback: pypdf / PyPDF2 annotation parsing
    if not links:
        try:
            try:
                from pypdf import PdfReader
            except ImportError:
                from PyPDF2 import PdfReader

            reader = PdfReader(path)
            for page in reader.pages:
                annots = page.get("/Annots", [])
                for annot in annots:
                    try:
                        obj = annot.get_object()
                        if obj.get("/Subtype") == "/Link":
                            uri = obj.get("/A", {}).get("/URI", "")
                            if uri:
                                links.append(uri if isinstance(uri, str) else uri.decode())
                    except Exception:
                        pass
        except Exception:
            pass

    return links

def extract_docx_links(path):
    """Extract all URLs embedded as hyperlinks in a DOCX."""
    links = []
    try:
        from docx import Document
        doc = Document(path)
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype.lower():
                target = rel._target
                if isinstance(target, str) and target.startswith("http"):
                    links.append(target)
    except Exception:
        pass
    return links

def extract_links(path):
    ext = os.path.splitext(path)[1].lower()
    if   ext == ".pdf":  return extract_pdf_links(path)
    elif ext == ".docx": return extract_docx_links(path)
    print(path)
    return []  # TXT files have no embedded links

# ── 3. Classify links into social profiles ────────────────────────────────────

LINK_PATTERNS = [
    ("linkedin",  r"linkedin\.com"),
    ("github",    r"github\.com"),
    ("medium",    r"medium\.com"),
    ("leetcode",  r"leetcode\.com"),
    ("twitter",   r"twitter\.com|x\.com"),
    ("kaggle",    r"kaggle\.com"),
    ("portfolio", r""),   # catch-all for personal websites
]

def classify_links(raw_links, text):
    """Map extracted URLs to named social slots, fall back to regex on text."""
    socials = {
        "email":     "",
        "linkedin":  "",
        "github":    "",
        "medium":    "",
        "leetcode":  "",
        "twitter":   "",
        "kaggle":    "",
        "website":   "",
    }

    def ensure_https(url):
        return url if url.startswith("http") else "https://" + url

    for url in raw_links:
        url = url.strip().rstrip("/")
        if not url:
            continue
        matched = False
        for slot, pattern in LINK_PATTERNS[:-1]:  # skip catch-all
            if pattern and re.search(pattern, url, re.I):
                if not socials[slot]:
                    socials[slot] = ensure_https(url)
                matched = True
                break
        if not matched and not socials["website"]:
            socials["website"] = ensure_https(url)

    # Fallback to regex on plain text for anything still missing
    if not socials["email"]:
        m = re.search(r"[\w.+\-]+@[\w\-]+\.[\w.]+", text)
        socials["email"] = m.group() if m else ""

    if not socials["linkedin"]:
        m = re.search(r"linkedin\.com/in/[\w\-]+", text, re.I)
        socials["linkedin"] = ("https://" + m.group()) if m else ""

    if not socials["github"]:
        m = re.search(r"github\.com/[\w\-]+", text, re.I)
        socials["github"] = ("https://" + m.group()) if m else ""

    if not socials["medium"]:
        m = re.search(r"medium\.com/@?[\w\-]+", text, re.I)
        socials["medium"] = ("https://" + m.group()) if m else ""

    if not socials["leetcode"]:
        m = re.search(r"leetcode\.com/[\w\-]+", text, re.I)
        socials["leetcode"] = ("https://" + m.group()) if m else ""

    # Remove empty keys so data.js stays clean
    return {k: v for k, v in socials.items() if v}

# ── 4. Split resume into named sections ──────────────────────────────────────

SECTION_HEADERS = [
    "experience", "work experience", "employment", "professional experience",
    "work history", "career history",
    "projects", "personal projects", "key projects", "notable projects",
    "skills", "technical skills", "core competencies", "technologies",
    "core technical skills", "key skills", "technical expertise",
    "tools & technologies", "tools and technologies", "tech stack",
    "skills & technologies", "skills and technologies",
    "education", "academic background",
    "certifications", "certificates", "courses",
    "summary", "profile", "objective", "about", "professional summary",
    "contact", "links",
]

def split_sections(text):
    lines = text.splitlines()
    sections = {}
    current = "header"
    buf = []
    for line in lines:
        stripped = line.strip()
        normalized = stripped.lower().rstrip(":").rstrip("-").strip()
        if normalized in SECTION_HEADERS and len(stripped) < 60:
            sections[current] = "\n".join(buf).strip()
            current = normalized
            buf = []
        else:
            buf.append(line)
    sections[current] = "\n".join(buf).strip()
    return sections

# ── 5. Content parsers ────────────────────────────────────────────────────────

def parse_name(header_text):
    for line in header_text.splitlines():
        line = line.strip()
        if line and not re.search(r"@|http|linkedin|github|phone|\d{3}", line, re.I):
            if len(line.split()) <= 5:
                return line
    return "Your Name"

def parse_title(header_text):
    title_keywords = ["engineer", "analyst", "scientist", "developer", "architect",
                      "manager", "lead", "consultant", "specialist", "designer"]
    for line in header_text.splitlines()[:8]:
        line = line.strip()
        # Strip everything after separators like · | - –
        for sep in [" · ", " | ", " - ", " – ", " — "]:
            if sep in line:
                line = line.split(sep)[0].strip()
                break
        if any(k in line.lower() for k in title_keywords) and len(line) < 80:
            return line
    return "Data Engineer"

def parse_location(header_text):
    for line in header_text.splitlines()[:5]:
        m = re.search(r"[A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z\s]+", line)
        if m and len(m.group()) < 60:
            return m.group().strip()
    return ""

def parse_skills(skills_text):
    if not skills_text:
        return []
    results = []
    lines = [l.strip() for l in skills_text.splitlines() if l.strip()]
    current_cat, current_tags = None, []
    for line in lines:
        colon_match = re.match(r"^([A-Za-z &/]+):\s*(.+)$", line)
        if colon_match:
            if current_cat and current_tags:
                results.append({"category": current_cat, "tags": current_tags})
            current_cat = colon_match.group(1).strip()
            current_tags = [t.strip() for t in re.split(r"[,|•·]", colon_match.group(2)) if t.strip()]
        elif current_cat:
            current_tags.extend([t.strip() for t in re.split(r"[,|•·]", line) if t.strip()])
        else:
            tags = [t.strip() for t in re.split(r"[,|•·\n]", line) if t.strip()]
            if tags:
                results.append({"category": "Technical Skills", "tags": tags})
    if current_cat and current_tags:
        results.append({"category": current_cat, "tags": current_tags})
    return results

DATE_RE = re.compile(
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|"
    r"April|June|July|August|September|October|November|December)[\w\s,]*\d{4}"
    r"\s*[-–—]\s*(?:Present|\d{4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\w\s,]*\d{4}))",
    re.I,
)

def is_job_header(line):
    """A job header line has ' | ' separating role from company and is not a bullet."""
    return " | " in line and not re.match(r"^[▸•\-\*►]", line)

def parse_experience(exp_text):
    if not exp_text:
        return []

    # Split into per-job blocks using "Role | Company" as delimiter
    raw_lines = exp_text.splitlines()
    blocks = []
    current = []
    for line in raw_lines:
        if is_job_header(line.strip()) and current:
            blocks.append(current)
            current = [line]
        else:
            current.append(line)
    if current:
        blocks.append(current)

    # If no pipe-separated headers found, fall back to blank-line splitting
    if len(blocks) <= 1:
        blocks = [b.splitlines() for b in re.split(r"\n{2,}", exp_text.strip())]

    entries = []
    for block in blocks:
        lines = [l.strip() for l in block if l.strip()]
        if not lines:
            continue

        role, company, period, bullets = "", "", "", []

        # First line: "Role | Company"  or just "Role"
        if is_job_header(lines[0]):
            parts = lines[0].split(" | ", 1)
            role    = parts[0].strip()
            company = parts[1].strip()
        else:
            role = lines[0]

        for line in lines[1:]:
            dm = DATE_RE.search(line)
            if dm and not period:
                period = dm.group().strip()
            elif re.match(r"^[▸•\-\*►]", line):
                cleaned = re.sub(r"^[▸•\-\*►]\s*", "", line)
                bullets.append(cleaned)
            elif not company and not period and not re.match(r"^[▸•\-\*►]", line):
                company = line

        if role:
            entries.append({"role": role, "company": company, "period": period, "bullets": bullets[:6]})
    return entries

def parse_projects(proj_text):
    if not proj_text:
        return []
    projects = []
    blocks = re.split(r"\n{2,}", proj_text.strip())
    for block in blocks:
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        title = lines[0]
        desc_lines, tags = [], []
        for line in lines[1:]:
            if re.match(r"^[A-Za-z0-9 ,|·•/+#]+$", line) and len(line) < 80 and "," in line:
                tags = [t.strip() for t in re.split(r"[,|•]", line) if t.strip()]
            elif re.match(r"^[•\-\*▸►]", line):
                desc_lines.append(re.sub(r"^[•\-\*▸►]\s*", "", line))
            else:
                desc_lines.append(line)
        # Extract any inline URLs
        github = ""
        demo   = ""
        for url in re.findall(r"https?://[\w./\-]+", block):
            if "github.com" in url and not github:
                github = url
            elif not demo:
                demo = url
        projects.append({"title": title, "desc": " ".join(desc_lines), "tags": tags, "github": github, "demo": demo})
    return projects

def parse_summary(summary_text):
    if not summary_text:
        return []
    paras = [p.strip() for p in re.split(r"\n{2,}", summary_text) if p.strip()]
    return paras[:3]

# ── 6. Assemble PORTFOLIO dict ────────────────────────────────────────────────

def build_portfolio(text, links):
    sections   = split_sections(text)
    header     = sections.get("header", text[:500])
    skills_key = next((k for k in sections if "skill" in k or "technolog" in k or "competenc" in k or "tool" in k or "stack" in k), None)
    exp_key    = next((k for k in sections if "experience" in k or "employment" in k or "history" in k), None)
    proj_key   = next((k for k in sections if "project" in k), None)
    summ_key   = next((k for k in sections if "summary" in k or "profile" in k or "about" in k or "objective" in k), None)

    name     = parse_name(header)
    title    = parse_title(header)
    location = parse_location(header)
    socials  = classify_links(links, text)

    skills     = parse_skills(sections.get(skills_key, ""))  if skills_key else []
    experience = parse_experience(sections.get(exp_key, "")) if exp_key    else []
    projects   = parse_projects(sections.get(proj_key, ""))  if proj_key   else []
    bio_paras  = parse_summary(sections.get(summ_key, ""))   if summ_key   else []

    if not bio_paras:
        bio_paras = [
            f"Hi! I'm a <strong>{title}</strong> with a passion for building robust data infrastructure.",
            "I work across the modern data stack — from ingestion and transformation to orchestration and delivery.",
        ]

    return {
        "name":     name,
        "title":    title,
        "tagline":  f"I'm a {title} passionate about building scalable data systems and turning raw data into business value.",
        "location": location,
        "socials":  socials,
        "about": {
            "bio": bio_paras,
            "highlights": [
                {"icon": "⚙️", "title": "Data Pipelines",       "desc": "ETL/ELT pipelines that move and transform data reliably at scale."},
                {"icon": "🏗️", "title": "Data Architecture",    "desc": "Designing lakes, warehouses, and lakehouses for analytical workloads."},
                {"icon": "📊", "title": "Analytics Engineering", "desc": "Transforming raw data into clean, trusted datasets for BI and reporting."},
            ],
        },
        "skills":     skills,
        "projects":   projects,
        "experience": experience,
    }

# ── 7. Write data.js ──────────────────────────────────────────────────────────

TEMPLATE = """\
// ============================================================
//  PORTFOLIO DATA — auto-generated from resume
//  Edit manually OR re-run:
//    python update_portfolio.py resume.pdf
// ============================================================

const PORTFOLIO = {content};
"""

def write_data_js(data):
    content = json.dumps(data, indent=2, ensure_ascii=False)
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data.js")
    with open(out, "w", encoding="utf-8") as f:
        f.write(TEMPLATE.replace("{content}", content))
    print(f"[OK] data.js written -> {out}")

# ── Main ──────────────────────────────────────────────────────────────────────

def find_resume():
    """Auto-detect a resume file in the same folder as this script."""
    folder = os.path.dirname(os.path.abspath(__file__))
    for ext in (".pdf", ".docx", ".txt"):
        for f in os.listdir(folder):
            if f.lower().endswith(ext) and "resume" in f.lower() or f.lower().endswith(ext) and "cv" in f.lower():
                return os.path.join(folder, f)
    # No obvious match — pick the first pdf/docx/txt that isn't this script
    for ext in (".pdf", ".docx", ".txt"):
        for f in os.listdir(folder):
            if f.lower().endswith(ext):
                return os.path.join(folder, f)
    return None

def main():
    # ── Resolve file path ────────────────────────────────────────────────────
    if len(sys.argv) >= 2:
        path = sys.argv[1]
    else:
        print("No file argument given — searching for a resume in this folder...")
        path = find_resume()
        if path:
            print(f"  Found: {path}")
        else:
            print("\nNo resume file found in the portfolio folder.")
            print("Put your resume here:  D:\\Claude\\portfolio\\resume.pdf")
            print("Then run:              python update_portfolio.py resume.pdf")
            sys.exit(1)

    if not os.path.exists(path):
        sys.exit(f"File not found: {path}\nMake sure the file is in D:\\Claude\\portfolio\\")

    # ── Extract text ─────────────────────────────────────────────────────────
    print(f"\n[1/4] Reading resume: {path}")
    text = read_resume(path)
    if not text.strip():
        sys.exit("Could not extract any text from the file. Is it a scanned image PDF? Try a text-based PDF or .docx instead.")
    print(f"      {len(text)} characters extracted")
    print(f"      Preview: {text[:120].strip()!r}...")

    # ── Extract links ─────────────────────────────────────────────────────────
    print(f"\n[2/4] Extracting embedded hyperlinks...")
    links = extract_links(path)
    if links:
        for lnk in links:
            print(f"      {lnk}")
    else:
        print("      None found — will fall back to regex on text")

    # ── Build portfolio data ──────────────────────────────────────────────────
    print(f"\n[3/4] Parsing resume sections...")
    sections = split_sections(text)
    print(f"      Sections detected: {list(sections.keys())}")

    data = build_portfolio(text, links)

    print(f"\n      Name:     {data['name']}")
    print(f"      Title:    {data['title']}")
    print(f"      Location: {data['location']}")
    print(f"      Skills:   {len(data['skills'])} categories")
    print(f"      Exp:      {len(data['experience'])} roles")
    print(f"      Projects: {len(data['projects'])} projects")
    print(f"      Socials:")
    for k, v in data["socials"].items():
        print(f"               {k:12} {v}")

    # ── Write data.js ─────────────────────────────────────────────────────────
    print(f"\n[4/4] Writing data.js...")
    write_data_js(data)
    print("\nDone! Open index.html in your browser to preview.")
    print("Tip: open data.js and tweak anything the parser missed.")

if __name__ == "__main__":
    main()
